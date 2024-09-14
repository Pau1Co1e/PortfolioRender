from fpdf import FPDF
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.text.run import Run
from fpdf.enums import XPos, YPos


class ResumeMaker:
    def __init__(self):
        self.pdf = FPDF()


# Function to sanitize text for FPDF
def sanitize_text(text):
    replacements = {
        "\u2013": "-",  # en-dash
        "\u2014": "-",  # em-dash
        "\u2018": "'",  # left single quotation mark
        "\u2019": "'",  # right single quotation mark
        "\u201C": '"',  # left double quotation mark
        "\u201D": '"',  # right double quotation mark
        "\u2026": "...",  # ellipsis
        "\u00A9": "(C)",  # copyright symbol
        "\u00AE": "(R)",  # registered trademark symbol
        "\u2122": "(TM)"  # trademark symbol
    }

    for unicode_char, replacement in replacements.items():
        text = text.replace(unicode_char, replacement)

    return text


# Function to create PDF resume (government or non-government formatted)
def create_pdf_resume(filename, is_government_format=True):
    pdf = FPDF()
    pdf.add_page()

    # Add name and contact info
    pdf.set_font('Helvetica', 'B', 16)
    pdf.cell(200, 10, 'Paul Coleman', new_x=XPos.LMARGIN, new_y=YPos.NEXT, align='C')

    pdf.set_font('Helvetica', '', 12)
    pdf.cell(200, 10, 'Email: engineering.paul@icloud.com', new_x=XPos.LMARGIN, new_y=YPos.NEXT, align='C')

    # Add hyperlinks for portfolio and GitHub
    pdf.set_font('Helvetica', 'U', 12)
    pdf.set_text_color(0, 0, 255)
    pdf.cell(200, 10, 'Portfolio Website', new_x=XPos.LMARGIN, new_y=YPos.NEXT, align='C', link='https://portfoliorender-p89i.onrender.com')
    pdf.cell(200, 10, 'GitHub', new_x=XPos.LMARGIN, new_y=YPos.NEXT, align='C', link='https://pau1co1e.github.io')

    if is_government_format:
        pdf.set_font('Helvetica', '', 12)
        pdf.set_text_color(0, 0, 0)
        pdf.cell(200, 10, 'Date Available to Begin Work: Immediately', new_x=XPos.LMARGIN, new_y=YPos.NEXT, align='C')
    pdf.set_font('Helvetica', '', 12)
    pdf.set_text_color(0, 0, 0)
    pdf.ln(10)

    # Summary section
    pdf.set_font('Helvetica', 'B', 14)
    pdf.cell(0, 10, 'SUMMARY', new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    pdf.set_font('Helvetica', '', 12)
    if is_government_format is False:
        pdf.multi_cell(190, 10, sanitize_text("Master’s student in Computer Science specializing in Artificial "
                                              "Intelligence and Machine Learning, with experience in full-stack "
                                              "development using the Microsoft tech stack. Skilled in developing "
                                              "AI-driven solutions for fintech and cybersecurity, leveraging a strong "
                                              "background in software development and data analysis. Adept at working "
                                              "with cloud infrastructure, containerization, and deployment of machine "
                                              "learning models. Proficient in Python, Flask, and Docker, "
                                              "with hands-on experience in designing, developing, and migrating web "
                                              "applications to AWS. Committed to utilizing technology to solve "
                                              "complex problems in dynamic environments."))
    if is_government_format:
        pdf.set_font('Helvetica', '', 12)
        pdf.multi_cell(190, 10, sanitize_text(
            "Versatile Software Engineer with expertise in Artificial Intelligence and Machine Learning, applied to "
            "cybersecurity, finance, data analysis, and cloud computing. Proven in developing AI models and machine "
            "learning algorithms, enhancing system performance and ensuring data security. Seeking an IT or Software "
            "Engineering role to improve system reliability and efficiency."
        ))

    pdf.ln(5)

    # Professional Experience section
    pdf.set_font('Helvetica', 'B', 14)
    pdf.cell(0, 10, 'PROFESSIONAL EXPERIENCE', new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    # Full Stack Software Engineer
    pdf.set_font('Helvetica', 'B', 12)
    pdf.cell(0, 10, 'Full Stack Software Engineer', new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.set_font('Helvetica', '', 12)
    if is_government_format is False:
        pdf.cell(0, 10, 'ResNexus | Salem, UT', new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    if is_government_format:
        pdf.cell(0, 10, 'ResNexus, Salem, UT | Sept 2023 - Dec 2023', new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.cell(0, 10, 'Salary: $70,000 per year | Hours per Week: 40 | Full-Time', new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.cell(0, 10, 'Supervisor: Kaycee Gibson, (801) 671-8760', new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    pdf.multi_cell(0, 10, sanitize_text(
        "- Developed backend APIs, frontend UI/UX, and database solutions, improving data integrity.\n"
        "- Worked with Microsoft tech stack, version control (Git), and deployment systems.\n"
        "- Enhanced hospitality software, ensuring system reliability and testing."
    ))

    pdf.ln(5)

    # Software Engineer Intern
    pdf.set_font('Helvetica', 'B', 12)
    pdf.cell(0, 10, 'Software Engineer Intern', new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.set_font('Helvetica', '', 12)

    if is_government_format is False:
        pdf.cell(0, 10, 'Utah Valley University College of Engineering', new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    if is_government_format:
        pdf.cell(0, 10, 'Utah Valley University College of Engineering | Jan 2021 - Dec 2021', new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.cell(0, 10, 'Salary: $20 per hour | Hours per Week: 30 | Part-Time', new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.cell(0, 10, 'Supervisor: Kazem Sohraby, (605) 786-5856', new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    pdf.multi_cell(0, 10, sanitize_text(
        "- Led a team enhancing robotics systems with AI for object recognition and NLP.\n"
        "- Implemented network security for robotics, ensuring data protection.\n"
        "- Showcased robotics projects at STEM fairs."
    ))

    pdf.ln(5)

    # Computer Science Assistant Teacher
    pdf.set_font('Helvetica', 'B', 12)
    pdf.cell(0, 10, 'Computer Science Assistant Teacher', new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.set_font('Helvetica', '', 12)

    if is_government_format is False:
        pdf.cell(0, 10, 'Utah Valley University Department of Computer Science', new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    if is_government_format:
        pdf.cell(0, 10, 'Utah Valley University Department of Computer Science | Jan 2020 - Dec 2020', new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.cell(0, 10, 'Salary: $15 per hour | Hours per Week: 25 | Part-Time', new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.cell(0, 10, 'Supervisor: Terry Hill, (801) 863-8218', new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    pdf.multi_cell(0, 10, sanitize_text(
        "- Teaching assistant for two undergraduate junior computer science courses; C++ and Python.\n"
        "- Co-designed curriculum, documentation, and coding standards.\n"
        "- Graded assignments and provided feedback and tutoring."
    ))

    pdf.ln(5)

    # Technical Support Specialist
    pdf.set_font('Helvetica', 'B', 12)
    pdf.cell(0, 10, 'Technical Support Specialist', new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.set_font('Helvetica', '', 12)
    if is_government_format is False:
        pdf.cell(0, 10, 'Dell Inc., Clearfield, UT', new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    if is_government_format:
        pdf.cell(0, 10, 'Dell Inc., Clearfield, UT | Jan 2007 - Jan 2008', new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.cell(0, 10, 'Salary: $13 per hour | Hours per Week: 40 | Full-Time', new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.cell(0, 10, 'Supervisor: Joshua Anderson, (877) 854-4965', new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    pdf.multi_cell(0, 10, sanitize_text(
        "- Provided technical support for customers of Dell XPS systems, focusing on troubleshooting hardware and "
        "software issues.\n"
        "- Received CompTIA A+ Hardware Certification."
    ))

    pdf.ln(5)

    # Education Section
    pdf.set_font('Helvetica', 'B', 14)
    pdf.cell(0, 10, 'EDUCATION', new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    pdf.set_font('Helvetica', '', 12)
    pdf.cell(0, 10, 'Master of Science in Computer Science', new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    if is_government_format:
        pdf.cell(0, 10, 'Utah Valley University | Expected: Aug 2025 | GPA: 3.27', new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    if is_government_format is False:
        pdf.cell(0, 10, 'Utah Valley University | Expected: Aug 2025', new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    pdf.ln(5)

    pdf.cell(0, 10, 'Bachelor of Science in Computer Science', new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    if is_government_format:
        pdf.cell(0, 10, 'Utah Valley University | Aug 2022 | GPA: 3.26', new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    if is_government_format is False:
        pdf.cell(0, 10, 'Utah Valley University | Aug 2022', new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    pdf.ln(5)

    # Certifications & Skills
    pdf.set_font('Helvetica', 'B', 14)
    pdf.cell(0, 10, 'CERTIFICATIONS & SKILLS', new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    pdf.set_font('Helvetica', '', 12)
    pdf.multi_cell(0, 10, sanitize_text(
        "- Programmer Certification | Utah Valley University | 2020\n"
        "- Deans List | 2020, 2021, 2024"
    ))
    pdf.ln(5)
    pdf.multi_cell(190, 10, sanitize_text(
        "- AI and ML: Model development, deep fake detection, Natural Language "
        "Processing (NLP), computer vision, autonomous systems, object and facial recognition, AI-driven "
        "cybersecurity.\n"
        "- Web Development: Flask, C#, JavaScript, Git, UI/UX.\n"
        "- Project Management: Agile, Scrum, leadership, communication.\n"
        "- Systems Architecture: Scalable systems, cloud computing (AWS, GCP).\n"
        "- Programming: Python, JavaScript, Java, SQL, HTML5/CSS3."
    ))

    pdf.ln(5)

    if is_government_format:
        # Additional Experience section
        pdf.set_font('Helvetica', 'B', 14)
        pdf.cell(0, 10, 'VOLUNTEER EXPERIENCE / COMMUNITY SERVICE', new_x=XPos.LMARGIN, new_y=YPos.NEXT)

        pdf.set_font('Helvetica', 'B', 12)
        pdf.cell(0, 10, 'Norther Utah Habitat for Humanity | Clothing Drive | June 2014', new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.set_font('Helvetica', '', 12)
        pdf.multi_cell(0, 10, sanitize_text(
            "Assisted in organized donation event for Habitat for Humanity at Utah State University, providing "
            "clothing for families in need."
        ))

    pdf.ln(5)

    # Save the PDF to file
    pdf.output(filename)
    print(f"PDF saved as {filename}")


def create_docx_resume(filename, is_government_format=True):
    doc = Document()

    # Add name and contact info
    doc.add_heading('Paul Coleman', level=0)
    doc.add_paragraph('Email: engineering.paul@icloud.com')

    # Website and GitHub with Alias Hyperlinks at the top
    contact_paragraph = doc.add_paragraph()
    add_hyperlink(contact_paragraph, 'Portfolio Website', 'https://portfoliorender-p89i.onrender.com')
    contact_paragraph.add_run(' | ')  # Adding separator between links
    add_hyperlink(contact_paragraph, 'GitHub', 'https://pau1co1e.github.io')

    doc.add_paragraph('Date Available to Begin Work: Immediately')

    # Summary section
    doc.add_heading('SUMMARY', level=1)
    if is_government_format:
        doc.add_paragraph(
            "Versatile Software Engineer with expertise in Artificial Intelligence and Machine Learning, applied to "
            "cybersecurity, finance, data analysis, and cloud computing. Proven in developing AI models and machine "
            "learning algorithms, enhancing system performance and ensuring data security. Seeking an IT or Software "
            "Engineering role to improve system reliability and efficiency."
        )
    if is_government_format is False:
        doc.add_paragraph("Master’s student in Computer Science specializing in Artificial Intelligence and Machine "
                          "Learning, with experience in full-stack development using the Microsoft tech stack. "
                          "Skilled in developing AI-driven solutions for fintech and cybersecurity, leveraging a "
                          "strong background in software development and data analysis. Adept at working with cloud "
                          "infrastructure, containerization, and deployment of machine learning models. Proficient in "
                          "Python, Flask, and Docker, with hands-on experience in designing, developing, "
                          "and migrating web applications to AWS. Committed to utilizing technology to solve complex "
                          "problems in dynamic environments.")

    # Professional Experience
    doc.add_heading('PROFESSIONAL EXPERIENCE', level=1)

    # Full Stack Software Engineer
    doc.add_heading('Full Stack Software Engineer', level=3)
    if is_government_format is False:
        doc.add_paragraph('ResNexus | Salem, UT')
    if is_government_format:
        doc.add_paragraph('ResNexus | Salem, UT | Sept 2023 - Dec 2023')
        doc.add_paragraph('Salary: $70,000 per year | Hours per Week: 40 | Full-Time')
        doc.add_paragraph('Supervisor: Kaycee Gibson, (801) 671-8760')
    doc.add_paragraph(
        "- Developed backend APIs, frontend UI/UX, and database solutions, improving data integrity.\n"
        "- Worked with Microsoft tech stack, version control (Git), and deployment systems.\n"
        "- Enhanced hospitality software, ensuring system reliability and testing."
    )

    # Software Engineer Intern
    doc.add_heading('Software Engineer Intern', level=3)
    if is_government_format is False:
        doc.add_paragraph('Utah Valley University College of Engineering | Orem, UT')
    if is_government_format:
        doc.add_paragraph('Utah Valley University College of Engineering | Jan 2021 - Dec 2021 | Orem, UT')
        doc.add_paragraph('Salary: $20 per hour | Hours per Week: 30 | Part-Time')
        doc.add_paragraph('Supervisor: Kazem Sohraby, (605) 786-5856')
    doc.add_paragraph(
        "- Led a team enhancing robotics systems with AI for object recognition and NLP.\n"
        "- Implemented network security for robotics, ensuring data protection.\n"
        "- Showcased robotics projects at STEM fairs."
    )

    # Computer Science Assistant Teacher
    doc.add_heading('Computer Science Assistant Teacher', level=3)
    if is_government_format is False:
        doc.add_paragraph('Utah Valley University Department of Computer Science | Orem, UT')
    if is_government_format:
        doc.add_paragraph('Utah Valley University | Jan 2020 - Dec 2020')
        doc.add_paragraph('Salary: $15 per hour | Hours per Week: 25 | Part-Time')
        doc.add_paragraph('Supervisor: Terry Hill, (801) 863-8218')

    doc.add_paragraph(
        "- Teaching assistant for two undergraduate junior computer science courses; C++ and Python.\n"
        "- Co-designed curriculum, documentation, and student requirements.\n"
        "- Graded student homework submissions and quizzes while providing feedback and tutoring."
    )

    # Technical Support Specialist
    doc.add_heading('Technical Support Specialist', level=3)
    doc.add_paragraph('Dell Inc. | Clearfield, UT')

    if is_government_format:
        doc.add_paragraph('Dell Inc. | Jan 2007 - Jan 2008 | Clearfield, UT')
        doc.add_paragraph('Salary: $13 per hour | Hours per Week: 40 | Full-Time')
        doc.add_paragraph('Supervisor: Joshua Anderson, (877) 854-4965')

    doc.add_paragraph(
        "- Provided technical support for Dell XPS systems, focusing on troubleshooting hardware/software issues.\n"
        "- Received CompTIA A+ Hardware Certification."
    )

    # Education Section
    doc.add_heading('EDUCATION', level=1)

    doc.add_heading('Master of Science in Computer Science', level=3)
    if is_government_format:
        doc.add_paragraph('Utah Valley University | Expected: Aug 2025 | GPA: 3.27')
    if is_government_format is False:
        doc.add_paragraph('Utah Valley University | Expected: Aug 2025')

    doc.add_heading('Bachelor of Science in Computer Science', level=3)
    if is_government_format:
        doc.add_paragraph('Utah Valley University | Aug 2022 | GPA: 3.26')
    if is_government_format is False:
        doc.add_paragraph('Utah Valley University | August 2022')

    # Certifications & Skills Section
    doc.add_heading('CERTIFICATIONS & SKILLS', level=1)

    # Subheadings for Skills
    doc.add_heading('Artificial Intelligence & Machine Learning', level=3)
    doc.add_paragraph(
        "Model development, deep fake detection, Natural Language Processing (NLP), computer vision, "
        "autonomous systems, object and facial recognition, AI-driven cybersecurity."
    )

    doc.add_heading('Web Development', level=3)
    doc.add_paragraph(
        "Flask, JavaScript, C#, Git, UI/UX design, responsive web development, version control."
    )

    doc.add_heading('Project Management', level=3)
    doc.add_paragraph(
        "Agile, Scrum, team leadership, communication, version control, development lifecycle management."
    )

    doc.add_heading('Systems Architecture', level=3)
    doc.add_paragraph(
        "Scalable systems, cloud computing (AWS, GCP), containerization, Docker, infrastructure optimization."
    )

    doc.add_heading('Programming Languages', level=3)
    doc.add_paragraph(
        "Python, TensorFlow, Java, SQL, Docker, C#, JavaScript, HTML5/CSS3."
    )

    # Additional Experience Section
    if is_government_format:
        doc.add_heading('VOLUNTEER EXPERIENCE / COMMUNITY SERVICE', level=1)
        doc.add_heading('Norther Utah Habitat for Humanity | Clothing Drive | June 2014', level=3)
        doc.add_paragraph(
            "Assisted in organized donation event for Habitat for Humanity at Utah State University, providing "
            "clothing for families in need."
        )

    # Save the DOCX file
    doc.save(filename)
    print(f"Word document saved as {filename}")


def add_hyperlink(paragraph, text, url):
    """
    A function that places a hyperlink within a paragraph.

    :param paragraph: The paragraph we are adding the hyperlink to.
    :param text: The display text for the link.
    :param url: The actual URL of the hyperlink.
    :return: None
    """
    # Create the hyperlink tag, and add needed values
    part = paragraph.part
    r_id = part.relate_to(url, "https://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
                          is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    # Create a run element and add the text to it
    run = OxmlElement('w:r')
    r_pr = OxmlElement('w:rPr')

    # Set the font color and underline
    r_style = OxmlElement('w:color')
    r_style.set(qn('w:val'), '0000FF')  # Blue color for hyperlink
    r_pr.append(r_style)

    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')  # Single underline
    r_pr.append(u)

    run.append(r_pr)
    run_text = OxmlElement('w:t')
    run_text.text = text
    run.append(run_text)

    hyperlink.append(run)
    paragraph._element.append(hyperlink)

    return hyperlink


# Main driver function to generate both formats
def main():

    # Government formatted resumes
    create_pdf_resume("government_resume.pdf", is_government_format=True)
    create_docx_resume("government_resume.docx", is_government_format=True)

    # Non-government formatted resumes
    create_pdf_resume("non_government_resume.pdf", is_government_format=False)
    create_docx_resume("non_government_resume.docx", is_government_format=False)


if __name__ == "__main__":
    main()
